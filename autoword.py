import os
import streamlit as st
from docx import Document
from docx.shared import Inches
from PIL import Image, ExifTags
import win32com.client

# Fonction pour corriger l'orientation des images
def correct_image_orientation(image_path):
    try:
        image = Image.open(image_path)
        exif = image._getexif()
        if exif:
            for tag, value in exif.items():
                tag_name = ExifTags.TAGS.get(tag, tag)
                if tag_name == "Orientation":
                    if value == 3:
                        image = image.rotate(180, expand=True)
                    elif value == 6:
                        image = image.rotate(270, expand=True)
                    elif value == 8:
                        image = image.rotate(90, expand=True)
                    image.save(image_path)
        return image_path
    except Exception as e:
        return None

# Fonction pour créer le document Word
def create_word_from_folder_structure(template_path, root_folder, output_path):
    try:
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Le modèle Word n'existe pas : {template_path}")
        
        doc = Document(template_path)

        def add_section_title(doc, section_name, level, is_first_section=False):
            paragraph = doc.add_paragraph(section_name)
            paragraph.style = f'Heading {min(level, 5)}'

        def process_folder(folder_path, level, is_first_section=False):
            folder_name = os.path.basename(folder_path)
            add_section_title(doc, folder_name, level, is_first_section)

            image_files = sorted([f for f in os.listdir(folder_path) if f.lower().endswith(('jpg', 'jpeg', 'png', 'gif'))])
            image_paths = list(dict.fromkeys([os.path.join(folder_path, f) for f in image_files]))  

            if image_paths:
                page_height = 9
                main_image_height = (4 / 5) * page_height
                main_image_width = main_image_height * 0.75
                small_image_height = (2 / 5) * main_image_height
                small_image_width = small_image_height * 0.75

                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Inches(main_image_width)
                table.columns[1].width = Inches(5)

                row = table.rows[0]
                cell_left = row.cells[0]
                paragraph = cell_left.paragraphs[0]
                run = paragraph.add_run()
                first_image_path = correct_image_orientation(image_paths.pop(0))
                run.add_picture(first_image_path, width=Inches(main_image_width), height=Inches(main_image_height))

                cell_right = row.cells[1]
                paragraph_right = cell_right.paragraphs[0]
                max_columns = 3
                col_count = 0

                for image in image_paths:
                    image_path = correct_image_orientation(image)
                    run = paragraph_right.add_run()
                    run.add_picture(image_path, width=Inches(small_image_width), height=Inches(small_image_height))
                    run.add_text("  ")
                    col_count += 1

                    if col_count >= max_columns:
                        paragraph_right.add_run("\n")
                        col_count = 0

                doc.add_page_break()

            first_subfolder = True
            for item in sorted(os.listdir(folder_path)):
                item_path = os.path.join(folder_path, item)
                if os.path.isdir(item_path):
                    process_folder(item_path, level + 1, is_first_section=first_subfolder)
                    first_subfolder = False

        process_folder(root_folder, 1, is_first_section=True)
        doc.save(output_path)
        update_table_of_contents(output_path)

    except Exception as e:
        print(f"❌ Erreur lors de la création du document : {e}")

def update_table_of_contents(doc_path):
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)
        doc.TablesOfContents(1).Update()
        word.Selection.EndKey(Unit=6)
        word.Selection.InsertBreak(3)
        doc.Save()
        doc.Close()
        word.Quit()
    except Exception as e:
        print(f"⚠️ Erreur lors de la mise à jour du sommaire : {e}")

# Streamlit interface
st.title("Création automatique de document Word à partir d'un dossier")
st.write("Ce programme parcourt un dossier et ses sous-dossiers pour créer un rapport avec des images.")

template_path = st.file_uploader("Choisissez un modèle Word", type=["docx"])
root_folder = st.text_input("Entrez le chemin du dossier contenant les images (root_folder)")
output_path = st.text_input("Entrez le chemin de sortie du document Word")

if st.button("Créer le document"):
    if template_path and root_folder and output_path:
        # Sauvegarder le modèle Word temporairement
        with open("template.docx", "wb") as f:
            f.write(template_path.getvalue())

        create_word_from_folder_structure("template.docx", root_folder, output_path)
        st.success(f"Le document a été créé avec succès : {output_path}")
    else:
        st.error("Veuillez remplir tous les champs.")
